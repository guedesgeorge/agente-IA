from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import Optional
import uvicorn, os, io, re
from datetime import datetime
from rag_engine import LexRAGEngine
from auth import (
    carregar_usuarios, salvar_usuarios, hash_senha, verificar_senha,
    criar_token, verificar_token, gerar_token_reset, verificar_token_reset,
    usar_token_reset, enviar_email_recuperacao,
    salvar_historico, carregar_historico, carregar_sessao, deletar_sessao
)

security = HTTPBearer(auto_error=False)

app = FastAPI(title="LexAI API", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
rag = LexRAGEngine()


def get_usuario_atual(credentials: HTTPAuthorizationCredentials = None):
    if not credentials:
        return None
    return verificar_token(credentials.credentials)


# ── MODELS ───────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    mode: str = "consulta"
    tipo_peca: Optional[str] = None
    municipio: Optional[str] = None
    session_id: Optional[str] = None
    use_knowledge_base: bool = True
    use_web_search: bool = True
    usuario: str = "advogado"
    history: list = []

class ChatResponse(BaseModel):
    answer: str
    sources: list[dict]
    web_searches: list[dict]
    session_id: str

class ExportRequest(BaseModel):
    content: str
    title: str = "Documento LexAI"
    municipio: str = ""

class LoginRequest(BaseModel):
    email: str
    senha: str

class SenhaRequest(BaseModel):
    email: str

class ResetSenhaRequest(BaseModel):
    token: str
    nova_senha: str

class AlterarSenhaRequest(BaseModel):
    senha_atual: str
    nova_senha: str

class SalvarHistoricoRequest(BaseModel):
    session_id: str
    titulo: str
    mensagens: list

class CriarUsuarioRequest(BaseModel):
    id: str
    nome: str
    email: str
    senha: str
    role: str = "advogado"

class EditarUsuarioRequest(BaseModel):
    nome: str = None
    email: str = None
    senha: str = None
    role: str = None


# ── HELPERS ──────────────────────────────────────────────────────────────────

def checar_admin(credentials):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    from auth import buscar_usuario_por_id
    u = buscar_usuario_por_id(user_id)
    if not u or u.get("role") != "admin":
        raise HTTPException(403, "Apenas administradores podem fazer isso")
    return user_id


# ── ROTAS PRINCIPAIS ─────────────────────────────────────────────────────────

@app.get("/")
def health():
    return {"status": "ok", "docs_indexed": rag.count_documents(), "web_search": True}

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    municipio: Optional[str] = None,
    tipo_peca: Optional[str] = None,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    allowed = [".pdf", ".docx", ".txt", ".doc"]
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(400, f"Formato não suportado: {ext}")
    content = await file.read()
    user_id = get_usuario_atual(credentials) if credentials else "publico"
    result = rag.ingest_document(content, file.filename, ext, municipio, tipo_peca, user_id=user_id)
    return {"status": "indexed", "filename": file.filename,
            "chunks": result["chunks"], "doc_id": result["doc_id"],
            "municipio": municipio or "geral", "user_id": user_id}

@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    import uuid
    session_id = req.session_id or str(uuid.uuid4())
    answer, sources, web_searches = await rag.answer(
        question=req.message,
        mode=req.mode,
        tipo_peca=req.tipo_peca,
        municipio=req.municipio,
        history=req.history,
        use_kb=req.use_knowledge_base,
        use_web=req.use_web_search
    )
    return ChatResponse(answer=answer, sources=sources,
                        web_searches=web_searches, session_id=session_id)

@app.get("/documents")
def list_documents(municipio: Optional[str] = None):
    return rag.list_documents(municipio)

@app.get("/documents/meus")
async def listar_documentos_usuario(
    municipio: Optional[str] = None,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    return rag.list_documents(municipio, user_id=user_id)

@app.delete("/documents/{doc_id}/usuario")
async def deletar_documento_usuario(
    doc_id: str,
    municipio: Optional[str] = None,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    docs = rag.list_documents(municipio, user_id=user_id)
    ids = [d.get("doc_id") for d in docs]
    if doc_id not in ids:
        raise HTTPException(403, "Documento não encontrado ou sem permissão")
    rag.delete_document(doc_id, municipio)
    return {"status": "deleted", "doc_id": doc_id}

@app.get("/tipos-peca")
def list_tipos():
    return rag.list_tipos_peca()

@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str, municipio: Optional[str] = None):
    rag.delete_document(doc_id, municipio)
    return {"status": "deleted", "doc_id": doc_id}


# ── EXPORT DOCX ──────────────────────────────────────────────────────────────

@app.post("/export-docx")
def export_docx(req: ExportRequest):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)

        if req.municipio:
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(f"MUNICÍPIO DE {req.municipio.upper()}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1F, 0x3F, 0x6B)

        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t.add_run(req.title.upper())
        tr.bold = True
        tr.font.size = Pt(13)
        tr.font.color.rgb = RGBColor(0x1F, 0x3F, 0x6B)
        doc.add_paragraph()

        for line in req.content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('#'):
                text = line.lstrip('#').strip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1F, 0x3F, 0x6B)
            elif len(line) < 60 and line.upper() == line and len(line) > 3:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1F, 0x3F, 0x6B)
            elif line.startswith('-') or line.startswith('•'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line.lstrip('-•').strip()).font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    if i % 2 == 1:
                        run.bold = True

        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(f"Gerado por LexAI · {datetime.now().strftime('%d/%m/%Y')}").font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_title = re.sub(r'[^\w\s-]', '', req.title).strip().replace(' ', '_')
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_title}.docx"}
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao gerar DOCX: {str(e)}")


# ── AUTENTICAÇÃO ──────────────────────────────────────────────────────────────

@app.post("/auth/login")
async def login(req: LoginRequest):
    from auth import buscar_usuario_por_email
    usuario = buscar_usuario_por_email(req.email)
    if not usuario:
        raise HTTPException(401, "Email ou senha incorretos")
    if not usuario.get("senha_hash"):
        raise HTTPException(401, "Senha ainda não definida. Use 'Esqueci minha senha' para criar.")
    if not verificar_senha(req.senha, usuario["senha_hash"]):
        raise HTTPException(401, "Email ou senha incorretos")
    token = criar_token(usuario["id"])
    return {
        "token": token,
        "usuario": {
            "id": usuario["id"],
            "nome": usuario["nome"],
            "email": usuario["email"],
            "role": usuario.get("role", "advogado")
        }
    }

@app.post("/auth/esqueci-senha")
async def esqueci_senha(req: SenhaRequest):
    from auth import buscar_usuario_por_email
    usuario = buscar_usuario_por_email(req.email)
    if usuario:
        token = gerar_token_reset(req.email)
        if token:
            await enviar_email_recuperacao(req.email, token, usuario["nome"])
    return {"ok": True, "msg": "Se o email estiver cadastrado, você receberá as instruções."}

@app.post("/auth/reset-senha")
async def reset_senha(req: ResetSenhaRequest):
    if len(req.nova_senha) < 6:
        raise HTTPException(400, "Senha deve ter pelo menos 6 caracteres")
    ok = usar_token_reset(req.token, req.nova_senha)
    if not ok:
        raise HTTPException(400, "Token inválido ou expirado")
    return {"ok": True, "msg": "Senha redefinida com sucesso!"}

@app.post("/auth/alterar-senha")
async def alterar_senha(req: AlterarSenhaRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    users = carregar_usuarios()
    usuario = users.get(user_id)
    if not usuario:
        raise HTTPException(404, "Usuário não encontrado")
    if usuario.get("senha_hash") and not verificar_senha(req.senha_atual, usuario["senha_hash"]):
        raise HTTPException(401, "Senha atual incorreta")
    if len(req.nova_senha) < 6:
        raise HTTPException(400, "Senha deve ter pelo menos 6 caracteres")
    users[user_id]["senha_hash"] = hash_senha(req.nova_senha)
    salvar_usuarios(users)
    return {"ok": True, "msg": "Senha alterada com sucesso!"}

@app.post("/auth/cadastrar-email")
async def cadastrar_email(email: str, credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    users = carregar_usuarios()
    if user_id not in users:
        raise HTTPException(404, "Usuário não encontrado")
    for uid, u in users.items():
        if uid != user_id and u.get("email", "").lower() == email.lower():
            raise HTTPException(400, "Email já cadastrado por outro usuário")
    users[user_id]["email"] = email
    salvar_usuarios(users)
    return {"ok": True, "msg": "Email cadastrado com sucesso!"}

@app.get("/auth/me")
async def get_me(credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    from auth import buscar_usuario_por_id
    usuario = buscar_usuario_por_id(user_id)
    if not usuario:
        raise HTTPException(404, "Usuário não encontrado")
    return {
        "id": usuario["id"],
        "nome": usuario["nome"],
        "email": usuario.get("email", ""),
        "role": usuario.get("role", "advogado")
    }


# ── HISTÓRICO ────────────────────────────────────────────────────────────────

@app.post("/historico/salvar")
async def salvar_hist(req: SalvarHistoricoRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    salvar_historico(user_id, req.session_id, req.titulo, req.mensagens)
    return {"ok": True}

@app.get("/historico")
async def listar_hist(credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    return carregar_historico(user_id)

@app.get("/historico/{session_id}")
async def get_sessao(session_id: str, credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    sessao = carregar_sessao(user_id, session_id)
    if not sessao:
        raise HTTPException(404, "Conversa não encontrada")
    return sessao

@app.delete("/historico/{session_id}")
async def deletar_hist(session_id: str, credentials: HTTPAuthorizationCredentials = Depends(security)):
    user_id = get_usuario_atual(credentials)
    if not user_id:
        raise HTTPException(401, "Não autenticado")
    deletar_sessao(user_id, session_id)
    return {"ok": True}


# ── ADMIN ─────────────────────────────────────────────────────────────────────

@app.get("/admin/usuarios")
async def listar_usuarios(credentials: HTTPAuthorizationCredentials = Depends(security)):
    checar_admin(credentials)
    users = carregar_usuarios()
    return [{"id": u["id"], "nome": u["nome"], "email": u.get("email", ""), "role": u.get("role", "")}
            for u in users.values()]

@app.post("/admin/usuarios")
async def criar_usuario(req: CriarUsuarioRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    checar_admin(credentials)
    users = carregar_usuarios()
    if req.id in users:
        raise HTTPException(400, "ID já existe")
    for u in users.values():
        if u.get("email", "").lower() == req.email.lower():
            raise HTTPException(400, "Email já cadastrado")
    if len(req.senha) < 6:
        raise HTTPException(400, "Senha deve ter mínimo 6 caracteres")
    users[req.id] = {
        "id": req.id,
        "nome": req.nome,
        "email": req.email,
        "senha_hash": hash_senha(req.senha),
        "role": req.role,
        "criado_em": datetime.now().isoformat()
    }
    salvar_usuarios(users)
    return {"ok": True, "msg": f"Usuário {req.nome} criado com sucesso"}

@app.put("/admin/usuarios/{user_id}")
async def editar_usuario(user_id: str, req: EditarUsuarioRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    checar_admin(credentials)
    users = carregar_usuarios()
    if user_id not in users:
        raise HTTPException(404, "Usuário não encontrado")
    if req.nome: users[user_id]["nome"] = req.nome
    if req.email:
        for uid, u in users.items():
            if uid != user_id and u.get("email", "").lower() == req.email.lower():
                raise HTTPException(400, "Email já usado por outro usuário")
        users[user_id]["email"] = req.email
    if req.senha:
        if len(req.senha) < 6:
            raise HTTPException(400, "Senha mínimo 6 caracteres")
        users[user_id]["senha_hash"] = hash_senha(req.senha)
    if req.role: users[user_id]["role"] = req.role
    salvar_usuarios(users)
    return {"ok": True, "msg": "Usuário atualizado"}

@app.delete("/admin/usuarios/{user_id}")
async def deletar_usuario(user_id: str, credentials: HTTPAuthorizationCredentials = Depends(security)):
    admin_id = checar_admin(credentials)
    if user_id == admin_id:
        raise HTTPException(400, "Não pode deletar a si mesmo")
    users = carregar_usuarios()
    if user_id not in users:
        raise HTTPException(404, "Usuário não encontrado")
    del users[user_id]
    salvar_usuarios(users)
    return {"ok": True, "msg": "Usuário removido"}

@app.post("/admin/usuarios/{user_id}/resetar-senha")
async def admin_resetar_senha(user_id: str, nova_senha: str, credentials: HTTPAuthorizationCredentials = Depends(security)):
    checar_admin(credentials)
    users = carregar_usuarios()
    if user_id not in users:
        raise HTTPException(404, "Usuário não encontrado")
    if len(nova_senha) < 6:
        raise HTTPException(400, "Senha mínimo 6 caracteres")
    users[user_id]["senha_hash"] = hash_senha(nova_senha)
    salvar_usuarios(users)
    return {"ok": True, "msg": "Senha redefinida com sucesso"}


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
